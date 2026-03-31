"""数据导入导出服务"""

import json
import re
import zipfile
import io
import shutil
from datetime import datetime
from pathlib import Path
from typing import Optional

from src.core.config import get_settings


class ExportImportService:
    """数据导入导出服务 - 支持 JSON/Markdown/ZIP 格式"""

    def __init__(self):
        self.settings = get_settings()

    async def export_project(self, project_id: str, format: str = "json") -> dict:
        """
        导出项目全部数据
        
        Args:
            project_id: 项目ID
            format: 导出格式 ('json', 'markdown', 'zip')
        
        Returns:
            dict: 导出结果信息
        """
        project_dir = self.settings.home_dir / ".aiwriter" / "projects" / project_id
        
        if not project_dir.exists():
            return {
                "success": False,
                "error": f"项目目录不存在: {project_dir}",
            }

        try:
            if format == "json":
                return await self._export_as_json(project_id, project_dir)
            elif format == "zip":
                return await self._export_as_zip(project_id, project_dir)
            else:
                return {
                    "success": False,
                    "error": f"不支持的导出格式: {format}",
                }
        except Exception as e:
            return {
                "success": False,
                "error": f"导出失败: {str(e)}",
            }

    async def _export_as_json(self, project_id: str, project_dir: Path) -> dict:
        """导出为单个 JSON 文件"""
        export_data = {
            "project_id": project_id,
            "exported_at": datetime.now().isoformat(),
            "version": self.settings.app_version,
            "data": {},
        }

        # 收集所有数据文件
        data_files = [
            "project.json",
            "chapters.json", 
            "characters.json",
            "relationships.json",
            "outlines.json",
            "inspirations.json",
            "locations.json",
            "world_settings.json",
        ]

        for filename in data_files:
            file_path = project_dir / filename
            if file_path.exists():
                with open(file_path, "r", encoding="utf-8") as f:
                    export_data["data"][filename.replace(".json", "")] = json.load(f)

        return {
            "success": True,
            "format": "json",
            "data": export_data,
        }

    async def _export_as_zip(self, project_id: str, project_dir: Path) -> dict:
        """导出为 ZIP 文件"""
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            # 添加项目数据文件
            for file_path in project_dir.iterdir():
                if file_path.is_file():
                    with open(file_path, "rb") as f:
                        zf.writestr(f"{project_id}/{file_path.name}", f.read())

            # 添加项目元数据
            meta = {
                "project_id": project_id,
                "exported_at": datetime.now().isoformat(),
                "version": self.settings.app_version,
            }
            zf.writestr(f"{project_id}/export_meta.json", json.dumps(meta, ensure_ascii=False, indent=2))

        zip_buffer.seek(0)
        
        return {
            "success": True,
            "format": "zip",
            "data": zip_buffer.getvalue(),
            "filename": f"{project_id}_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
        }

    async def export_chapter_markdown(self, chapter_id: str) -> dict:
        """
        导出章节为 Markdown（包含角色标注）
        
        Args:
            chapter_id: 章节ID
        
        Returns:
            dict: 包含 markdown 内容的导出结果
        """
        try:
            # 获取章节数据
            chapter = await self._get_chapter(chapter_id)
            if not chapter:
                return {
                    "success": False,
                    "error": f"章节 {chapter_id} 未找到",
                }

            project_id = chapter.get("project_id")
            title = chapter.get("title", "无标题")
            content = chapter.get("content", "")
            characters = chapter.get("characters", [])

            # 构建 Markdown 内容
            md_lines = [f"# {title}\n"]
            
            if characters:
                md_lines.append("## 登场角色\n")
                for char in characters:
                    char_name = char.get("name", "未知")
                    char_desc = char.get("description", "")
                    md_lines.append(f"- **{char_name}**: {char_desc}")
                md_lines.append("\n")

            md_lines.append("## 正文\n")
            md_lines.append(content)

            # 如果有备注
            if chapter.get("notes"):
                md_lines.append("\n\n## 作者备注\n")
                md_lines.append(f"_{chapter['notes']}_")

            markdown_content = "\n".join(md_lines)

            return {
                "success": True,
                "format": "markdown",
                "data": {
                    "chapter_id": chapter_id,
                    "project_id": project_id,
                    "title": title,
                    "content": markdown_content,
                },
            }

        except Exception as e:
            return {
                "success": False,
                "error": f"导出 Markdown 失败: {str(e)}",
            }

    async def _get_chapter(self, chapter_id: str) -> Optional[dict]:
        """获取章节数据"""
        # 遍历项目目录查找章节
        projects_base = self.settings.home_dir / ".aiwriter" / "projects"
        
        for project_dir in projects_base.iterdir():
            chapters_file = project_dir / "chapters.json"
            if chapters_file.exists():
                with open(chapters_file, "r", encoding="utf-8") as f:
                    chapters = json.load(f)
                    for ch in chapters if isinstance(chapters, list) else chapters.get("items", []):
                        if ch.get("id") == chapter_id:
                            ch["project_id"] = project_dir.name
                            return ch
        return None

    async def import_project(self, file_data: bytes, format: str = "json") -> dict:
        """
        导入项目数据
        
        Args:
            file_data: 文件数据
            format: 导入格式 ('json', 'zip')
        
        Returns:
            dict: 导入结果
        """
        try:
            if format == "json":
                return await self._import_from_json(file_data)
            elif format == "zip":
                return await self._import_from_zip(file_data)
            else:
                return {
                    "success": False,
                    "error": f"不支持的导入格式: {format}",
                }
        except Exception as e:
            return {
                "success": False,
                "error": f"导入失败: {str(e)}",
            }

    async def _import_from_json(self, file_data: bytes) -> dict:
        """从 JSON 导入"""
        try:
            data = json.loads(file_data.decode("utf-8"))
        except UnicodeDecodeError:
            data = json.loads(file_data)

        project_id = data.get("project_id")
        if not project_id:
            return {
                "success": False,
                "error": "导入数据缺少 project_id",
            }

        project_dir = self.settings.home_dir / ".aiwriter" / "projects" / project_id
        project_dir.mkdir(parents=True, exist_ok=True)

        imported_files = []
        project_data = data.get("data", {})

        for filename, content in project_data.items():
            file_path = project_dir / f"{filename}.json"
            with open(file_path, "w", encoding="utf-8") as f:
                json.dump(content, f, ensure_ascii=False, indent=2)
            imported_files.append(f"{filename}.json")

        return {
            "success": True,
            "project_id": project_id,
            "imported_files": imported_files,
            "imported_at": datetime.now().isoformat(),
        }

    async def _import_from_zip(self, file_data: bytes) -> dict:
        """从 ZIP 导入"""
        zip_buffer = io.BytesIO(file_data)
        project_id = None
        imported_files = []

        with zipfile.ZipFile(zip_buffer, "r") as zf:
            # 获取项目 ID
            names = zf.namelist()
            for name in names:
                if name.endswith("/"):
                    # 目录名即项目ID
                    project_id = name.rstrip("/")
                    break

            if not project_id:
                # 尝试从 export_meta.json 获取
                if "export_meta.json" in names:
                    meta = json.loads(zf.read("export_meta.json"))
                    project_id = meta.get("project_id")

            if not project_id:
                return {
                    "success": False,
                    "error": "无法确定项目ID",
                }

            project_dir = self.settings.home_dir / ".aiwriter" / "projects" / project_id
            project_dir.mkdir(parents=True, exist_ok=True)

            # 解压文件
            for name in names:
                if name.endswith("/") or name == "export_meta.json":
                    continue

                content = zf.read(name)
                filename = Path(name).name
                file_path = project_dir / filename
                
                with open(file_path, "wb") as f:
                    f.write(content)
                imported_files.append(filename)

        return {
            "success": True,
            "project_id": project_id,
            "imported_files": imported_files,
            "imported_at": datetime.now().isoformat(),
        }

    async def import_chapters_from_folder(self, project_id: str, folder_path: str) -> dict:
        """
        从文件夹批量导入章节（Markdown 文件）
        
        Args:
            project_id: 目标项目ID
            folder_path: 包含 Markdown 文件的文件夹路径
        
        Returns:
            dict: 导入结果
        """
        try:
            folder = Path(folder_path)
            if not folder.exists() or not folder.is_dir():
                return {
                    "success": False,
                    "error": f"文件夹不存在: {folder_path}",
                }

            project_dir = self.settings.home_dir / ".aiwriter" / "projects" / project_id
            project_dir.mkdir(parents=True, exist_ok=True)

            # 读取现有章节
            chapters_file = project_dir / "chapters.json"
            existing_chapters = []
            if chapters_file.exists():
                with open(chapters_file, "r", encoding="utf-8") as f:
                    existing_chapters = json.load(f)

            if isinstance(existing_chapters, dict):
                existing_chapters = existing_chapters.get("items", [])

            existing_ids = {ch.get("id") for ch in existing_chapters}

            imported_chapters = []
            new_chapter_id = f"{project_id}_ch{len(existing_chapters) + 1}"

            # 扫描 Markdown 文件
            md_files = sorted(folder.glob("*.md")) + sorted(folder.glob("*.markdown"))
            
            for i, md_file in enumerate(md_files):
                with open(md_file, "r", encoding="utf-8") as f:
                    content = f.read()

                # 解析 Markdown
                title = md_file.stem  # 默认使用文件名作为标题
                body_content = content
                characters = []

                # 尝试解析标题（# 标题）
                if content.startswith("# "):
                    lines = content.split("\n", 1)
                    title = lines[0][2:].strip()
                    body_content = lines[1] if len(lines) > 1 else ""

                # 尝试解析角色（## 登场角色）
                if "## 登场角色" in body_content:
                    parts = body_content.split("## 登场角色")
                    body_content = parts[0]
                    char_section = parts[1].split("##")[0] if len(parts) > 1 else ""
                    
                    for line in char_section.strip().split("\n"):
                        if line.startswith("- **"):
                            # 格式: - **角色名**: 描述
                            line = line[4:]
                            name_part, desc = line.split("**:", 1)
                            characters.append({
                                "id": f"{new_chapter_id}_char{len(characters) + 1}",
                                "name": name_part.strip(),
                                "description": desc.strip(),
                            })

                # 去除正文中的 ## 作者备注 等
                if "## 作者备注" in body_content:
                    body_content = body_content.split("## 作者备注")[0]

                chapter = {
                    "id": new_chapter_id,
                    "project_id": project_id,
                    "title": title,
                    "content": body_content.strip(),
                    "characters": characters,
                    "order": len(existing_chapters) + i + 1,
                    "created_at": datetime.now().isoformat(),
                    "updated_at": datetime.now().isoformat(),
                }
                imported_chapters.append(chapter)
                existing_chapters.append(chapter)
                existing_ids.add(new_chapter_id)
                new_chapter_id = f"{project_id}_ch{len(existing_chapters) + 1}"

            # 保存更新后的章节
            with open(chapters_file, "w", encoding="utf-8") as f:
                json.dump(existing_chapters, f, ensure_ascii=False, indent=2)

            return {
                "success": True,
                "project_id": project_id,
                "imported_count": len(imported_chapters),
                "imported_chapters": [
                    {"id": ch["id"], "title": ch["title"]} 
                    for ch in imported_chapters
                ],
                "imported_at": datetime.now().isoformat(),
            }

        except Exception as e:
            return {
                "success": False,
                "error": f"批量导入章节失败: {str(e)}",
            }

    async def export_world_settings(self, project_id: str, format: str = "md") -> dict:
        """
        导出项目设定集
        
        Args:
            project_id: 项目ID
            format: 导出格式 ('md' 或 'json')
        
        Returns:
            dict: 导出结果
        """
        try:
            project_dir = self.settings.home_dir / ".aiwriter" / "projects" / project_id
            
            if not project_dir.exists():
                return {
                    "success": False,
                    "error": f"项目目录不存在: {project_dir}",
                }
            
            # 读取设定数据
            world_settings_file = project_dir / "world_settings.json"
            if not world_settings_file.exists():
                return {
                    "success": False,
                    "error": "项目没有设定集数据",
                }
            
            with open(world_settings_file, "r", encoding="utf-8") as f:
                world_settings = json.load(f)
            
            if isinstance(world_settings, dict):
                settings_list = world_settings.get("items", [world_settings])
            else:
                settings_list = world_settings
            
            # 按类别分组
            categorized = {}
            for setting in settings_list:
                category = setting.get("category", "other")
                if category not in categorized:
                    categorized[category] = []
                categorized[category].append(setting)
            
            if format == "json":
                return {
                    "success": True,
                    "format": "json",
                    "data": {
                        "project_id": project_id,
                        "exported_at": datetime.now().isoformat(),
                        "categories": categorized,
                        "total_count": len(settings_list),
                    },
                }
            else:
                # Markdown 格式
                md_lines = [f"# {project_id} 世界设定集\n"]
                md_lines.append(f"> 导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                md_lines.append(f"> 共 {len(settings_list)} 项设定\n")
                md_lines.append("")
                
                # 按类别输出
                category_names = {
                    "character": "人物",
                    "location": "地点",
                    "item": "物品",
                    "organization": "组织",
                    "concept": "概念",
                }
                
                for category, settings in categorized.items():
                    category_name = category_names.get(category, category)
                    md_lines.append(f"## {category_name}\n")
                    
                    for setting in settings:
                        name = setting.get("name", "未命名")
                        content = setting.get("content", "")
                        md_lines.append(f"### {name}\n")
                        if content:
                            md_lines.append(f"{content}\n")
                        else:
                            md_lines.append("_暂无内容_\n")
                        md_lines.append("")
                    
                    md_lines.append("---\n")
                
                return {
                    "success": True,
                    "format": "markdown",
                    "data": {
                        "project_id": project_id,
                        "content": "\n".join(md_lines),
                        "total_count": len(settings_list),
                    },
                }
                
        except Exception as e:
            return {
                "success": False,
                "error": f"导出设定集失败: {str(e)}",
            }

    async def export_chapter_as_text(self, chapter_id: str) -> dict:
        """
        导出章节为纯文本 TXT（去除 Markdown 格式）
        
        Args:
            chapter_id: 章节ID
        
        Returns:
            dict: 包含纯文本内容的导出结果
        """
        import re
        
        try:
            chapter = await self._get_chapter(chapter_id)
            if not chapter:
                return {
                    "success": False,
                    "error": f"章节 {chapter_id} 未找到",
                }
            
            project_id = chapter.get("project_id")
            title = chapter.get("title", "无标题")
            content = chapter.get("content", "")
            
            # 去除 Markdown 格式
            text_content = content
            
            # 去除 # 标题标记
            text_content = re.sub(r'^#{1,6}\s+', '', text_content, flags=re.MULTILINE)
            # 去除 ** 粗体
            text_content = re.sub(r'\*\*(.*?)\*\*', r'\1', text_content)
            # 去除 * 斜体
            text_content = re.sub(r'\*(.*?)\*', r'\1', text_content)
            # 去除 ~~ 删除线
            text_content = re.sub(r'~~(.*?)~~', r'\1', text_content)
            # 去除 ` 行内代码
            text_content = re.sub(r'`(.*?)`', r'\1', text_content)
            # 去除 [链接](url) 转为链接文字
            text_content = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text_content)
            # 去除图片语法
            text_content = re.sub(r'!\[.*?\]\(.*?\)', '', text_content)
            # 去除列表标记
            text_content = re.sub(r'^[\s]*[-*+]\s+', '', text_content, flags=re.MULTILINE)
            text_content = re.sub(r'^[\s]*\d+\.\s+', '', text_content, flags=re.MULTILINE)
            # 去除引用标记
            text_content = re.sub(r'^>\s+', '', text_content, flags=re.MULTILINE)
            # 去除水平线
            text_content = re.sub(r'^[-*_]{3,}$', '', text_content, flags=re.MULTILINE)
            # 去除多余空行
            text_content = re.sub(r'\n{3,}', '\n\n', text_content)
            
            return {
                "success": True,
                "format": "txt",
                "data": {
                    "chapter_id": chapter_id,
                    "project_id": project_id,
                    "title": title,
                    "content": text_content.strip(),
                },
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"导出纯文本失败: {str(e)}",
            }

    async def export_chapter_as_json(self, chapter_id: str) -> dict:
        """
        导出章节为结构化 JSON 格式
        
        Args:
            chapter_id: 章节ID
        
        Returns:
            dict: 包含结构化 JSON 的导出结果
        """
        try:
            chapter = await self._get_chapter(chapter_id)
            if not chapter:
                return {
                    "success": False,
                    "error": f"章节 {chapter_id} 未找到",
                }
            
            return {
                "success": True,
                "format": "json",
                "data": {
                    "chapter_id": chapter_id,
                    "project_id": chapter.get("project_id"),
                    "title": chapter.get("title", "无标题"),
                    "content": chapter.get("content", ""),
                    "summary": chapter.get("summary", ""),
                    "notes": chapter.get("notes", ""),
                    "word_count": len(chapter.get("content", "")),
                    "characters": chapter.get("characters", []),
                    "order": chapter.get("order", 0),
                    "created_at": chapter.get("created_at", ""),
                    "updated_at": chapter.get("updated_at", ""),
                },
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"导出 JSON 失败: {str(e)}",
            }

    async def export_chapter_as_docx(self, chapter_id: str) -> bytes:
        """
        导出章节为 Word 文档（使用 python-docx）
        
        Args:
            chapter_id: 章节ID
        
        Returns:
            bytes: .docx 文件的二进制内容
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            chapter = await self._get_chapter(chapter_id)
            if not chapter:
                raise ValueError(f"章节 {chapter_id} 未找到")
            
            title = chapter.get("title", "无标题")
            content = chapter.get("content", "")
            
            # 创建 Word 文档
            doc = Document()
            
            # 设置标题
            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加角色信息（如果有）
            characters = chapter.get("characters", [])
            if characters:
                doc.add_heading("登场角色", level=2)
                for char in characters:
                    char_name = char.get("name", "未知")
                    char_desc = char.get("description", "")
                    p = doc.add_paragraph()
                    p.add_run(f"• {char_name}").bold = True
                    if char_desc:
                        p.add_run(f": {char_desc}")
                doc.add_paragraph()
            
            # 添加正文
            doc.add_heading("正文", level=2)
            
            # 解析内容（简单的段落分割）
            paragraphs = content.split("\n\n")
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue
                
                # 跳过已经是标题的
                if para_text.startswith("#"):
                    continue
                
                # 添加段落
                p = doc.add_paragraph(para_text)
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.first_line_indent = Pt(24)  # 首行缩进
            
            # 添加作者备注（如果有）
            if chapter.get("notes"):
                doc.add_paragraph()
                doc.add_heading("作者备注", level=2)
                notes_para = doc.add_paragraph(chapter["notes"])
                notes_para.paragraph_format.line_spacing = 1.5
            
            # 保存到字节流
            from io import BytesIO
            docx_stream = BytesIO()
            doc.save(docx_stream)
            docx_stream.seek(0)
            
            return docx_stream.getvalue()
            
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")
        except Exception as e:
            raise Exception(f"导出 Word 文档失败: {str(e)}")


# 全局服务实例
_export_import_service: Optional[ExportImportService] = None


def get_export_import_service() -> ExportImportService:
    """获取导入导出服务实例"""
    global _export_import_service
    if _export_import_service is None:
        _export_import_service = ExportImportService()
    return _export_import_service
